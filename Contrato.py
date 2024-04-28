from builtins import str
from docx2pdf import convert
import pandas as pd
import os
from docx import Document
from docx.shared import  Pt,Inches,Cm,Mm
import datetime
dfc=pd.read_excel('Contratos 03-12-20.xlsx')
dfd=pd.read_excel('Contratos 04-12-20.xlsx')
###########################################
########Ciclo Iterativo ###################
###########################################

def ContratoProfesional(Sexo, Nombres, Apellido_Paterno, Apellido_Materno, Nacionalidad, Direccion, Comuna, Rut,
                        Fecha_inicio, Fecha_Fin, Monto_Total,Rol):
    ###########################################
    ########DOCUMENTO WORD ####################
    ###########################################

    document = Document()
    ###########################################
    ########PARAMETROS DEL DOCUMENTO  #########
    ###########################################
    font = document.styles['Normal'].font
    font.name = 'Book Antiqua'
    sections = document.sections
    for section in sections:
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
    ###########################################
    ########CREACION DEL DOCUMENTO  ###########
    ###########################################

    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("header.png")

    # kh.add_picture('header.png')

    document.add_picture('logo.png')

    h = document.add_paragraph('')
    h.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS A HONORARIOS\n').bold = True
    h.alignment = 1
    font.size = Pt(10)
    h.add_run(
        '............................................................................................................................................................................................................................. ').bold = True
    p = document.add_paragraph(
        'En Temuco, a '+str(Fecha_inicio)+', entre la Fundación de Desarrollo Educacional y Tecnológico La Araucanía,Rut71.195.600-K, representada por su Director Ejecutivo don ISMAEL SIMON TOLOZA BRAVO, Cédula de Identidad Nº11.355.138-0, ambos domiciliados en calle Montevideo Nº0780 de esta ciudad, en adelante la “Fundación” y  '+Sexo+'  '+Nombres+'  '+Apellido_Paterno+'  '+Apellido_Materno+', de nacionalidad '+Nacionalidad+', con domicilio en '+Direccion+', de la ciudad de '+Comuna+', Cédula de Identidad N°'+Rut+', en adelante, el “Prestador de Servicios”, se ha convenido el siguiente contrato de prestación de servicios a honorarios: \n')
    font.size = Pt(8)
    p1 = document.add_paragraph('')
    p1.add_run('PRIMERO        :').bold = True
    p1.add_run(
        'En el contexto del contrato de prestación de servicios con  fecha 11 de Noviembre de 2020, suscrito entre La Agencia de Calidad de la Educación y la Fundación de Desarrollo Educacional y Tecnológico La Araucanía en unión temporal de proveedores con la Universidad de La Frontera, y aprobado según Resolución Exenta N°603del 23 de Noviembre de 2020. La Fundación contrata los servicios a honorarios del Prestador de Servicios, para que realice la siguiente prestación de servicio en el proyecto “Servicios asociados a la evaluación de conocimientos específicos y pedagógicos 2020, ID 721703­19­LR20.”')
    p2 = document.add_paragraph('')
    p2.add_run('SEGUNDO        :').bold = True
    p2.add_run('El rol a desempeñar es de '+Rol+'.')
    p3 = document.add_paragraph('')
    p3.add_run('TERCERO        :').bold = True
    p3.add_run(
        'El plazo para la realización de la prestación de servicios encomendada será del '+str(Fecha_inicio)+'hasta el '+str(Fecha_Fin))
    p4 = document.add_paragraph('')
    p4.add_run('CUARTO        :').bold = True
    p4.add_run(
        'Por la prestación de servicios efectivamente realizada, se pagará un monto total bruto de $'+str(Monto_Total)+'.-pagaderos contra presentación de boleta de honorarios respectiva, debiendo ser visada por el Jefe de Proyecto. El pago de los honorarios se efectuará según el calendario de entrega de Boletas de Honorarios y de pago establecido por la Fundación. De los honorarios pactados, la Fundación retendrá el porcentaje correspondiente al Impuesto a la Renta que se integrará en la Tesorería General de la República.')
    p5 = document.add_paragraph('')
    p5.add_run('QUINTO        :').bold = True
    p5.add_run('El Prestador de Servicios acepta el encargo y las condiciones precedentes.')
    p6 = document.add_paragraph('')
    p6.add_run('SEXTO        :').bold = True
    p6.add_run(
        'El presente contrato no obliga a la Fundación a mantener  la continuidad del servicio prestado o trabajo convenido, al cual podrá poner término en forma unilateral y sin mayor reclamo, bastando la comunicación  a la otra parte de forma personal o escrita, en el último caso emitida por carta certificada al domicilio fijado en el contrato, entendiéndose notificada al tercer día de despachada la carta, produciendo plenos efectos, desde dicho momento sin necesidad de fundamentar las razones de la decisión adoptada.')
    p7 = document.add_paragraph('')
    p7.add_run('SEPTIMO        :').bold = True
    p7.add_run(
        'El Prestador de Servicios deberá resguardar el carácter de confidencial de todo el material de aplicación según el Acuerdo de Confidencialidad firmado.')
    p8 = document.add_paragraph('')
    p8.add_run('OCTAVO        :').bold = True
    p8.add_run(
        'Las partes expresan que este contrato no constituye vínculo de dependencia entre el Prestador de Servicios y la Fundación contratante, ya que las labores del Prestador de Servicios son discontinuas, no son exclusivas y no se realizan bajo ningún vínculo de subordinación.')
    p9 = document.add_paragraph('')
    p9.add_run('NOVENO        :').bold = True
    p9.add_run('El costo  de  este  convenio  será  imputado  al Centro de Costo Nº20267“ECEP 2020”.')
    p10 = document.add_paragraph('')
    p10.add_run('DECIMO        :').bold = True
    p10.add_run(
        'Para todos los efectos legales de este contrato, las partes fijan su domicilio en la ciudad de Temuco.')
    p11 = document.add_paragraph('')
    p11.add_run('En comprobante, previa lectura y ratificación, las partes firman.  ').bold = True
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells[1].add_paragraph()
    r = hdr_cells0.add_run()
    r.add_picture('firma.png')
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '-----------------------------------------------------------\nEL PRESTADOR DE SERVICIOS'
    hdr_cells[1].text = '-----------------------------------------------------------\np. LA FUNDACIÓN'

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Montevideo 0780, Temuco, Chile, +56 45 2325900. www.fudeaufro.cl')
    run.add_picture("footer1.png")



    ###########################################
    ########CREACION DE DIRECTORIO   ##########
    ###########################################
    dir = Nombres + ' ' + Apellido_Paterno + ' ' + Apellido_Materno
    if not os.path.exists(dir):
        os.mkdir(str(dir))
        rutafinal=str(dir) + '\ ' + str(dir) + '.docx'
        document.save(rutafinal)



def ContratoRelator(Sexo, Nombre_Completo, Nacionalidad, Direccion, Comuna, Rut,
                        Fecha_inicio, Fecha_Fin, Monto_Total,Rol):
    ###########################################
    ########DOCUMENTO WORD ####################
    ###########################################

    document = Document()
    ###########################################
    ########PARAMETROS DEL DOCUMENTO  #########
    ###########################################
    font = document.styles['Normal'].font
    font.name = 'Book Antiqua'
    sections = document.sections
    for section in sections:
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
    ###########################################
    ########CREACION DEL DOCUMENTO  ###########
    ###########################################

    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("header.png")

    # kh.add_picture('header.png')

    document.add_picture('logo.png')

    h = document.add_paragraph('')
    h.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS A HONORARIOS\n').bold = True
    h.alignment = 1
    font.size = Pt(10)
    h.add_run(
        '............................................................................................................................................................................................................................. ').bold = True
    p = document.add_paragraph(
        'En Temuco, a '+str(Fecha_inicio)+', entre la Fundación de Desarrollo Educacional y Tecnológico La Araucanía, Rut 71.195.600-K, representada por su Director Ejecutivo don ISMAEL SIMON TOLOZA BRAVO, Cédula de Identidad Nº11.355.138-0, ambos domiciliados en calle Montevideo Nº0780 de esta ciudad, en adelante la “Fundación” y  '+Sexo+'  '+Nombre_Completo+', de nacionalidad '+Nacionalidad+', con domicilio en '+Direccion+', de la ciudad de '+Comuna+', Cédula de Identidad N°'+Rut+', en adelante, el “Prestador de Servicios”, se ha convenido el siguiente contrato de prestación de servicios a honorarios: \n')
    font.size = Pt(8)
    p1 = document.add_paragraph('')
    p1.add_run('PRIMERO        :').bold = True
    p1.add_run(
        'En el contexto del contrato de prestación de servicios con  fecha 11 de Noviembre de 2020, suscrito entre La Agencia de Calidad de la Educación y la Fundación de Desarrollo Educacional y Tecnológico La Araucanía en unión temporal de proveedores con la Universidad de La Frontera, y aprobado según Resolución Exenta N°603del 23 de Noviembre de 2020. La Fundación contrata los servicios a honorarios del Prestador de Servicios, para que realice la siguiente prestación de servicio en el proyecto “Servicios asociados a la evaluación de conocimientos específicos y pedagógicos 2020, ID 721703­19­LR20.”')
    p2 = document.add_paragraph('')
    p2.add_run('SEGUNDO        :').bold = True
    p2.add_run('El rol a desempeñar es de '+Rol+'.')
    p3 = document.add_paragraph('')
    p3.add_run('TERCERO        :').bold = True
    p3.add_run(
        'El plazo para la realización de la prestación de servicios encomendada será del '+str(Fecha_inicio)+'hasta el '+str(Fecha_Fin))
    p4 = document.add_paragraph('')
    p4.add_run('CUARTO        :').bold = True
    p4.add_run('Por el servicio profesional efectivamente realizado, se pagara un monto bruto variable, el cual corresponderá a cada proceso de capacitación que realice, de acuerdo al siguiente detalle: ')
    table4 = document.add_table(rows=2, cols=2)
    table4.alignment = 1
    hdr_cells0 = table4.rows[0].cells
    hdr_cells0[0].text='Proceso de Capacitación'
    hdr_cells0[1].text='Monto Bruto'
    hdr_cells = table4.rows[1].cells
    hdr_cells[0].text = Rol
    hdr_cells[1].text = Monto_Total
    p5 = document.add_paragraph('')
    p5.add_run('QUINTO        :').bold = True
    p5.add_run('El Prestador de Servicios acepta el encargo y las condiciones precedentes.')
    p6 = document.add_paragraph('')
    p6.add_run('SEXTO        :').bold = True
    p6.add_run(
        'El presente contrato no obliga a la Fundación a mantener  la continuidad del servicio prestado o trabajo convenido, al cual podrá poner término en forma unilateral y sin mayor reclamo, bastando la comunicación  a la otra parte de forma personal o escrita, en el último caso emitida por carta certificada al domicilio fijado en el contrato, entendiéndose notificada al tercer día de despachada la carta, produciendo plenos efectos, desde dicho momento sin necesidad de fundamentar las razones de la decisión adoptada.')
    p7 = document.add_paragraph('')
    p7.add_run('SEPTIMO        :').bold = True
    p7.add_run(
        'El Prestador de Servicios deberá resguardar el carácter de confidencial de todo el material de aplicación según el Acuerdo de Confidencialidad firmado.')
    p8 = document.add_paragraph('')
    p8.add_run('OCTAVO        :').bold = True
    p8.add_run(
        'Las partes expresan que este contrato no constituye vínculo de dependencia entre el Prestador de Servicios y la Fundación contratante, ya que las labores del Prestador de Servicios son discontinuas, no son exclusivas y no se realizan bajo ningún vínculo de subordinación.')
    p9 = document.add_paragraph('')
    p9.add_run('NOVENO        :').bold = True
    p9.add_run('El costo  de  este  convenio  será  imputado  al Centro de Costo Nº20267“ECEP 2020”.')
    p10 = document.add_paragraph('')
    p10.add_run('DECIMO        :').bold = True
    p10.add_run(
        'Para todos los efectos legales de este contrato, las partes fijan su domicilio en la ciudad de Temuco.')
    p11 = document.add_paragraph('')
    p11.add_run('En comprobante, previa lectura y ratificación, las partes firman.  ').bold = True
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells[1].add_paragraph()
    r = hdr_cells0.add_run()
    r.add_picture('firma.png')
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '-----------------------------------------------------------\nEL PRESTADOR DE SERVICIOS'
    hdr_cells[1].text = '-----------------------------------------------------------\np. LA FUNDACIÓN'

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Montevideo 0780, Temuco, Chile, +56 45 2325900. www.fudeaufro.cl')
    run.add_picture("footer1.png")
   

    ###########################################
    ########CREACION DE DIRECTORIO   ##########
    ###########################################
    dir = Nombre_Completo
    if not os.path.exists(dir):
        os.mkdir(str(dir))
        document.save(str(dir) +'\ '+ str(dir) + '.docx')
        #ruta relativa pdf


def ContratoJefeDeSede(Sexo, Nombre_Completo, Nacionalidad, Direccion, Comuna, Rut,
                    Fecha_inicio, Fecha_Fin, Monto_Total, Rol):
    ###########################################
    ########DOCUMENTO WORD ####################
    ###########################################

    document = Document()
    ###########################################
    ########PARAMETROS DEL DOCUMENTO  #########
    ###########################################
    font = document.styles['Normal'].font
    font.name = 'Book Antiqua'
    sections = document.sections
    for section in sections:
        section.page_height = Cm(35.56)
        section.page_width = Cm(21.59)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
    ###########################################
    ########CREACION DEL DOCUMENTO  ###########
    ###########################################

    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("header.png")

    # kh.add_picture('header.png')

    document.add_picture('logo.png')

    h = document.add_paragraph('')
    h.add_run('CONTRATO DE PRESTACIÓN DE SERVICIOS A HONORARIOS\n').bold = True
    h.alignment = 1
    font.size = Pt(10)
    h.add_run(
        '............................................................................................................................................................................................................................. ').bold = True
    p = document.add_paragraph(
        'En Temuco, a ' + str(
            Fecha_inicio) + ', entre la Fundación de Desarrollo Educacional y Tecnológico La Araucanía,Rut71.195.600-K, representada por su Director Ejecutivo don ISMAEL SIMON TOLOZA BRAVO, Cédula de Identidad Nº11.355.138-0, ambos domiciliados en calle Montevideo Nº0780 de esta ciudad, en adelante la “Fundación” y  ' + Sexo + '  ' + Nombre_Completo + ', de nacionalidad ' + Nacionalidad + ', con domicilio en ' + Direccion + ', de la ciudad de ' + Comuna + ', Cédula de Identidad N°' + Rut + ', en adelante, el “Prestador de Servicios”, se ha convenido el siguiente contrato de prestación de servicios a honorarios: \n')
    font.size = Pt(8)
    p1 = document.add_paragraph('')
    p1.add_run('PRIMERO        :').bold = True
    p1.add_run(
        'En el contexto del contrato de prestación de servicios con  fecha 11 de Noviembre de 2020, suscrito entre La Agencia de Calidad de la Educación y la Fundación de Desarrollo Educacional y Tecnológico La Araucanía en unión temporal de proveedores con la Universidad de La Frontera, y aprobado según Resolución Exenta N°603del 23 de Noviembre de 2020. La Fundación contrata los servicios a honorarios del Prestador de Servicios, para que realice la siguiente prestación de servicio en el proyecto “Servicios asociados a la evaluación de conocimientos específicos y pedagógicos 2020, ID 721703­19­LR20.”')
    p2 = document.add_paragraph('')
    p2.add_run('SEGUNDO        :').bold = True
    p2.add_run('El rol a desempeñar es de ' + Rol + '.')
    p3 = document.add_paragraph('')
    p3.add_run('TERCERO        :').bold = True
    p3.add_run(
        'El plazo para la realización de la prestación de servicios encomendada será del ' + str(
            Fecha_inicio) + 'hasta el ' + str(Fecha_Fin))
    p4 = document.add_paragraph('')
    p4.add_run('CUARTO        :').bold = True
    p4.add_run(
        'Por el servicio profesional efectivamente realizado, se pagara un monto bruto variable, el cual corresponderá a cada proceso de capacitación que realice, de acuerdo al siguiente detalle: ')
    table4 = document.add_table(rows=5, cols=2)
    table4.alignment = 1
    hdr_cells0 = table4.rows[0].cells
    hdr_cells0[0].text = 'Producto'
    hdr_cells0[1].text = 'Monto Bruto'
    hdr_cells = table4.rows[1].cells
    hdr_cells[0].text = Rol
    hdr_cells[1].text = Monto_Total[0]
    hdr_cells = table4.rows[2].cells
    hdr_cells[0].text = Rol
    hdr_cells[1].text = Monto_Total[1]
    hdr_cells = table4.rows[3].cells
    hdr_cells[0].text = Rol
    hdr_cells[1].text = Monto_Total[2]
    hdr_cells = table4.rows[4].cells
    hdr_cells[0].text = Rol
    hdr_cells[1].text = Monto_Total[3]
    p4.add_run('Pagaderos contra presentación de boleta de honorarios respectiva, debiendo ser visada por el Jefe de Proyecto. El pago de los honorarios se efectuará según el calendario de entrega de Boletas de Honorarios y de pago establecido por la Fundación. De los honorarios pactados, la Fundación retendrá el porcentaje correspondiente al Impuesto a la Renta que se integrará en la Tesorería General de la República.')
    p5 = document.add_paragraph('')
    p5.add_run('QUINTO        :').bold = True
    p5.add_run('El Prestador de Servicios acepta el encargo y las condiciones precedentes.')
    p6 = document.add_paragraph('')
    p6.add_run('SEXTO        :').bold = True
    p6.add_run(
        'El presente contrato no obliga a la Fundación a mantener  la continuidad del servicio prestado o trabajo convenido, al cual podrá poner término en forma unilateral y sin mayor reclamo, bastando la comunicación  a la otra parte de forma personal o escrita, en el último caso emitida por carta certificada al domicilio fijado en el contrato, entendiéndose notificada al tercer día de despachada la carta, produciendo plenos efectos, desde dicho momento sin necesidad de fundamentar las razones de la decisión adoptada.')
    p7 = document.add_paragraph('')
    p7.add_run('SEPTIMO        :').bold = True
    p7.add_run(
        'El Prestador de Servicios deberá resguardar el carácter de confidencial de todo el material de aplicación según el Acuerdo de Confidencialidad firmado.')
    p8 = document.add_paragraph('')
    p8.add_run('OCTAVO        :').bold = True
    p8.add_run(
        'Las partes expresan que este contrato no constituye vínculo de dependencia entre el Prestador de Servicios y la Fundación contratante, ya que las labores del Prestador de Servicios son discontinuas, no son exclusivas y no se realizan bajo ningún vínculo de subordinación.')
    p9 = document.add_paragraph('')
    p9.add_run('NOVENO        :').bold = True
    p9.add_run('El costo  de  este  convenio  será  imputado  al Centro de Costo Nº20267“ECEP 2020”.')
    p10 = document.add_paragraph('')
    p10.add_run('DECIMO        :').bold = True
    p10.add_run(
        'Para todos los efectos legales de este contrato, las partes fijan su domicilio en la ciudad de Temuco.')
    p11 = document.add_paragraph('')
    p11.add_run('En comprobante, previa lectura y ratificación, las partes firman.  ').bold = True
    table = document.add_table(rows=2, cols=2)
    table.alignment = 1
    hdr_cells0 = table.rows[0].cells[1].add_paragraph()
    r = hdr_cells0.add_run()
    r.add_picture('firma.png')
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = '-----------------------------------------------------------\nEL PRESTADOR DE SERVICIOS'
    hdr_cells[1].text = '-----------------------------------------------------------\np. LA FUNDACIÓN'

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run('Montevideo 0780, Temuco, Chile, +56 45 2325900. www.fudeaufro.cl')
    run.add_picture("footer1.png")


    ###########################################
    ########CREACION DE DIRECTORIO   ##########
    ###########################################
    dir = Nombre_Completo
    if not os.path.exists(dir):
        os.mkdir(str(dir))
        rutafinal=str(dir) + '\ ' + str(dir) + '.docx'
        document.save(rutafinal)
        # ruta relativa pdf
    os.mkdir('PDF')
    Carpeta= 'PDF'
    convert(str(dir) + '\ ' + str(dir) + '.docx','PDF/'+Nombre_Completo+'.pdf')


def Creacion_Contratos():
    for index, row in dfc.iterrows():
        Sexo, Nombres, Apellido_Paterno, Apellido_Materno, Nacionalidad, Direccion, Comuna, Rut, Fecha_inicio, Fecha_Fin, Monto_Total,Rol = dfc[['Sexo', 'Nombres', 'Apellido Paterno',
             'Apellido Materno', 'Nacionalidad', 'Dirección', 'Comuna', 'RUT', 'Desde', 'Hasta'
            , 'Monto','ROL']].iloc[index]
        ContratoProfesional(Sexo,Nombres,Apellido_Paterno,Apellido_Materno,Nacionalidad,Direccion,Comuna,Rut,Fecha_inicio,Fecha_Fin,Monto_Total,Rol)

#FILTRADO En base a Contratos Relator
dfd_filtrado=dfd.loc[dfd['ROL'].str.contains('Relator')]
dfd_filtrado.reset_index(drop=True,inplace=True)

def Creacion_Contratos_Relator():
    for index, row in dfd_filtrado.iterrows():
        Sexo,Nombre_Completo, Nacionalidad, Direccion, Comuna, Rut, Fecha_inicio, Fecha_Fin, Monto_Total,Rol = dfd_filtrado[['Unnamed: 8','Nombre Completo', 'Nacionalidad', 'Dirección', 'Ciudad', 'RUT', 'Desde', 'Hasta'
            , 'Monto Bruto','ROL']].iloc[index]
        ContratoRelator(Sexo,Nombre_Completo,Nacionalidad,Direccion,Comuna,Rut,Fecha_inicio,Fecha_Fin,Monto_Total,Rol)


dfd_filtrado2=dfd.loc[dfd['ROL'].str.contains('Jefe de Sede')]
dfd_filtrado.reset_index(drop=True,inplace=True)
def Creacion_Contratos_Jefe_Sede():
    for index, row in dfd_filtrado.iterrows():
        Sexo,Nombre_Completo, Nacionalidad, Direccion, Comuna, Rut, Fecha_inicio, Fecha_Fin, Monto_Total,Rol = dfd_filtrado[['Sexo','Nombre Completo', 'Nacionalidad', 'Dirección', 'Comuna', 'RUT', 'Desde', 'Hasta'
            , 'Monto Bruto','ROL']].iloc[index]
        if(Monto_Total == 'Según Tabla 2019.'):
            Monto_Total=['160.000 ','136.000','114.000','90000']
        ContratoRelator(Sexo,Nombre_Completo,Nacionalidad,Direccion,Comuna,Rut,Fecha_inicio,Fecha_Fin,Monto_Total,Rol)



#Creacion_Contratos()
#Creacion_Contratos_Relator()
print('antes de filtrar')
print(dfd)
dfd_filtrado2=dfd.loc[dfd['ROL'].str.contains('Jefe de Sede')]
dfd_filtrado.reset_index(drop=True,inplace=True)
# wrkite dataframe to excel file with no index
#dfd_filtrado2.to_excel("jefe.xls")